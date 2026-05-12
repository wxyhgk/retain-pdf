from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def set_section_page(section, *, width_pt: float, height_pt: float) -> None:
    section.page_width = Pt(width_pt)
    section.page_height = Pt(height_pt)
    section.left_margin = Pt(0)
    section.right_margin = Pt(0)
    section.top_margin = Pt(0)
    section.bottom_margin = Pt(0)
    section.header_distance = Pt(0)
    section.footer_distance = Pt(0)


def add_background_image(document: Document, image_path: Path, *, width_pt: float, height_pt: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Pt(width_pt), height=Pt(height_pt))
    send_last_inline_picture_behind_text(run)


def add_page_break(document: Document) -> None:
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def send_last_inline_picture_behind_text(run) -> None:
    drawings = run._r.xpath(".//w:drawing")
    if not drawings:
        return
    inline = drawings[-1].find(qn("wp:inline"))
    if inline is None:
        return
    inline.tag = qn("wp:anchor")
    inline.set("distT", "0")
    inline.set("distB", "0")
    inline.set("distL", "0")
    inline.set("distR", "0")
    inline.set("simplePos", "0")
    inline.set("relativeHeight", "0")
    inline.set("behindDoc", "1")
    inline.set("locked", "0")
    inline.set("layoutInCell", "1")
    inline.set("allowOverlap", "1")

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")

    pos_h = OxmlElement("wp:positionH")
    pos_h.set("relativeFrom", "page")
    pos_h_offset = OxmlElement("wp:posOffset")
    pos_h_offset.text = "0"
    pos_h.append(pos_h_offset)

    pos_v = OxmlElement("wp:positionV")
    pos_v.set("relativeFrom", "page")
    pos_v_offset = OxmlElement("wp:posOffset")
    pos_v_offset.text = "0"
    pos_v.append(pos_v_offset)

    inline.insert(0, simple_pos)
    inline.insert(1, pos_h)
    inline.insert(2, pos_v)
    extent = inline.find(qn("wp:extent"))
    wrap_none = OxmlElement("wp:wrapNone")
    if extent is not None:
        inline.insert(list(inline).index(extent) + 1, wrap_none)
    else:
        inline.insert(3, wrap_none)
